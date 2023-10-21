import requests
from docx import Document
from datetime import datetime
from docx.text.run import Run
from docx.shared import Pt
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.dml import MSO_THEME_COLOR_INDEX
import re
import docx
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from datetime import date

# Jira API base URL
print("Your Jira Cloud URL should look something like yourcompany.atlassian.com")
BASE_URL = input("Jira Cloud URL:")
BASE_URL = BASE_URL + "/rest/api/2"

# Jira credentials
USERNAME= input("Username: ")
print("--------------------------------------------------------")
print("Next you'll be asked for an API token.  If you don't have a token, here's how to set one up.")
print("1. In Jira, click your picture in the upper right and choose _manage account_")
print("2. In the profile page, choose the Security tab on the top.")
print("3. Under Security, go to Create and manage API tokens")
print("--------------------------------------------------------")
API_TOKEN = input("Please enter your Jira API Token: ")

# Create a session with basic authentication
session = requests.Session()
session.auth = (USERNAME, API_TOKEN)



def insertHR(paragraph):
    p = paragraph._p  # p is the <w:p> XML element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr,
        'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
        'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
        'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
        'w:pPrChange'
    )
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)

def add_hyperlink(paragraph, url, text, color, underline):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Add color if it is given
    if not color is None:
      c = docx.oxml.shared.OxmlElement('w:color')
      c.set(docx.oxml.shared.qn('w:val'), color)
      rPr.append(c)

    # Remove underlining if it is requested
    if not underline:
      u = docx.oxml.shared.OxmlElement('w:u')
      u.set(docx.oxml.shared.qn('w:val'), 'none')
      rPr.append(u)

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink

def get_issues(jql_query):
    # Define your JQL query to filter issues
    #jql_query = input("Enter the JQL query for the report: ")
    #"example: type in (epic)

    # Make a GET request to Jira API to fetch issues
    response = session.get(f"{BASE_URL}/search", params={"jql": jql_query})

    if response.status_code == 200:
        return response.json()["issues"]
    else:
        print(f"Failed to fetch issues: {response.status_code}")
        return []
    
def generate_report(issues,docx_filename):
    # Customize your report generation logic here
    doc=Document()

    currentdate=datetime.today()
    datestring=currentdate.strftime("%m/%d/%y")
    doc.add_heading(f"Jira Report - {datestring}", level=1)
    doc.add_heading(f"for {USERNAME}",level=2)
    p=doc.add_paragraph("")
    insertHR(p)
    for issue in issues:
        try:
            mailtext=f"{issue['fields']['assignee']['displayName']}"
        except:
            mailtext="Undefined"
        try:
            maillink=f"mailto://{issue['fields']['assignee']['emailAddress']}"
        except:
            maillink="https://teirpoint.atlassian.net"
        issuetext=f"{issue['key']}"
        weblink=f"https://tierpoint.atlassian.net/browse/{issuetext}"
        doc.add_heading(f"{issue['fields']['summary']}", level=3)
        p=doc.add_paragraph(f"{issue['fields']['priority']['name']} {issue['fields']['issuetype']['name']} \t Status: {issue['fields']['status']['name']} | Progress: {issue['fields']['progress']['progress']} | ETA:{issue['fields']['customfield_10023']} \t")
        hyperlink=add_hyperlink(p,weblink,issuetext,"FF8822",True)
        p=doc.add_paragraph(f"Updated: {issue['fields']['updated']} \t Assigned: ")
        hyperlink=add_hyperlink(p,maillink,mailtext,"FF8822",True)
        try:
            p=doc.add_paragraph(f"Latest Update: {issue['fields']['customfield_10078']}")
        except:
            p=doc.add_paragraph("No update")
        p.paragraph_format.left_indent=Inches(0.25)
        p=doc.add_paragraph("")
        insertHR(p)

    doc.add_heading(f"Query used in this report:",level=2)
    p=doc.add_paragraph(f"{jql_query}")
    p.paragraph_format.left_indent=Inches(0.25)
   
    doc.save(docx_filename)


if __name__ == "__main__":
    jql_query = input("Enter the JQL query for the report: ")
    issues = get_issues(jql_query)
    filename=input("Filename for the report (jira_report.docx): ") or "jira_report.docx"
    generate_report(issues,filename)